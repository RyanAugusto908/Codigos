
from docx import Document

doc = Document()
doc.add_paragraph("Este é um exemplo de manipulação de arquivos docx usando python-docx.")
doc.add_paragraph("Lembre-se de sempre salvar o arquivo após escrever nele.")
doc.save("exemplo.docx")

# Lendo arquivo .docx:
doc = Document("exemplo.docx")
print("Conteúdo do arquivo .docx:")
for paragrafo in doc.paragraphs:
    print(paragrafo.text)

# Usando 'with' para manipular arquivos de forma mais segura.
# Python-docx não utiliza 'with' diretamente como open(), mas o fluxo é seguro assim.
doc = Document()
doc.add_paragraph("Este é um exemplo de manipulação segura usando python-docx")
doc.save("exemplo_with.docx")

# Lendo o arquivo criado:
doc = Document("exemplo_with.docx")
print("\nConteúdo usando python-docx: ")
for paragrafo in doc.paragraphs:
    print(paragrafo.text)

    