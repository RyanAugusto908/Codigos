from docx import Document


def criar_documento():
    doc = Document()
    doc.add_paragraph("Teste do meu arquivo word, criando dados.")
    doc.add_paragraph("O código para criar parágrafo é 'doc.add_paragraph'")
    doc.save("Modificar_arquivo.docx")

def adicionar_dados(doc_path,nome,idade):
    
    #Pedir dados ao usuário:
    nome = str(input("Digite seu nome: "))
    idade = int(input("Sua idade: "))
    
    try:
        doc = Document(doc_path)
        doc.add_paragraph("Criação de dados:")
        doc.add_paragraph(f"Nome: {nome}")
        doc.add_paragraph(f"Idade: {idade}")
        doc.save(doc_path)
    except Exception as e:
        print(f"Ocorreu um erro ao adicionar dados {e}")

def ler_documento(doc_path):
    try:
        doc = Document(doc_path)
        print("Conteúdo do arquivo .docx:")
        for paragrafo in doc.paragraphs:
            print(paragrafo.text)
    except Exception as e:
        print(f"Ocorreu um erro ao ler o documento: {e}")

# Execução do código
criar_documento()
adicionar_dados("Modificar_arquivo.docx", "a", 1)
ler_documento("Modificar_arquivo.docx")