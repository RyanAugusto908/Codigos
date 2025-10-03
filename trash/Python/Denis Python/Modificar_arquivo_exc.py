from docx import Document

#Base Word:
doc = Document()
doc.add_paragraph("Teste do meu arquivo word, criando dados.")
doc.add_paragraph("O código para criar paragrafo é 'doc.add_paragraph'")
doc.save("Modificar_arquivo.docx")

#Criar dados e adicionar ao documento Word
doc = Document("Modificar_arquivo.docx")  # Abrir o documento existente
doc.add_paragraph("Criação de dados:")
doc.add_paragraph("Nome: Ryan")
doc.add_paragraph("Idade: 19")
doc.save("Modificar_arquivo.docx")  # Salvar as alterações

#Lendo arquivo:    
doc = Document("Modificar_arquivo.docx")
print("Conteúdo do arquivo .docx:")
for paragrafo in doc.paragraphs:
    print(paragrafo.text)


    
